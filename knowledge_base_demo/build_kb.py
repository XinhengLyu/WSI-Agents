"""
Knowledge base build script.
Supports two data sources:
  1. JSON structured medical knowledge documents (data/ directory)
  2. Word (.docx) medical book documents (book/ directory)

Usage:
  python build_kb.py
"""

import os
from pathlib import Path
from typing import List

from docx import Document
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
import json


# ─── Configuration ──────────────────────────────────────────────────────────

OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY", "")   # Set via environment variable

# Chroma vector store persistence directory (must match main project)
PERSIST_DIRECTORY = "medical_kb_structured"

# Data source directories (relative to this file)
JSON_DATA_DIR = "../data"          # Structured JSON knowledge documents
DOCX_DATA_DIR = "../book"          # WHO medical books (.docx files)


# ─── JSON Document Processing ────────────────────────────────────────────────

class MedicalDocumentProcessor:
    """Extracts structured text from JSON-format medical knowledge documents"""

    @staticmethod
    def extract_text(json_str: str) -> str:
        data = json.loads(json_str)
        text_parts = []

        if "titles" in data:
            titles = [data["titles"][lvl] for lvl in ("h1", "h2", "h3") if lvl in data["titles"]]
            text_parts.append(" > ".join(titles))

        if "sections" in data:
            for section in data["sections"]:
                text_parts.append(f"\n# {section['title']}")
                if "content" in section:
                    text_parts.append("\n".join(f"- {item}" for item in section["content"]))

        return "\n".join(text_parts)


def load_json_documents(base_dir: str) -> List[LangchainDocument]:
    """Recursively read all .json files under base_dir, return LangchainDocument list"""
    base_path = Path(base_dir)
    documents  = []

    for file_path in sorted(base_path.rglob("*.json")):
        if file_path.name.startswith("."):
            continue
        try:
            relative_path = file_path.relative_to(base_path)
            category      = relative_path.parts[0]
            content       = file_path.read_text(encoding="utf-8")
            text          = MedicalDocumentProcessor.extract_text(content)

            documents.append(LangchainDocument(
                page_content=text,
                metadata={
                    "file_name": file_path.name,
                    "category":  category,
                    "full_path": str(file_path),
                    "source":    str(relative_path),
                }
            ))
            print(f"  [JSON] Loaded: {relative_path}")
        except Exception as e:
            print(f"  [JSON] Skipped {file_path.name}: {e}")

    return documents


# ─── DOCX Document Processing ─────────────────────────────────────────────────

class DocxProcessor:
    """Splits Word documents into text chunks"""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""]
        )

    def process(self, file_path: str, category: str = "medical_books") -> List[LangchainDocument]:
        doc       = Document(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chunks    = self.text_splitter.split_text(full_text)

        return [
            LangchainDocument(
                page_content=chunk,
                metadata={
                    "file_name": Path(file_path).name,
                    "category":  category,
                    "full_path": str(file_path),
                    "chunk_id":  i,
                }
            )
            for i, chunk in enumerate(chunks)
        ]


def load_docx_documents(directory: str, category: str = "medical_books") -> List[LangchainDocument]:
    """Read all .docx/.doc files under directory"""
    processor = DocxProcessor()
    documents  = []

    for root, _, files in os.walk(directory):
        for name in files:
            if not name.endswith((".docx", ".doc")):
                continue
            path = os.path.join(root, name)
            try:
                docs = processor.process(path, category)
                documents.extend(docs)
                print(f"  [DOCX] Processed: {name}, generated {len(docs)} chunks")
            except Exception as e:
                print(f"  [DOCX] Skipped {name}: {e}")

    return documents


# ─── Knowledge Base Building ──────────────────────────────────────────────────

def build_knowledge_base(
    json_dir:   str = JSON_DATA_DIR,
    docx_dir:   str = DOCX_DATA_DIR,
    persist_dir: str = PERSIST_DIRECTORY,
) -> Chroma:
    """
    Build and persist the vector knowledge base.
    Indexes both JSON structured documents and DOCX medical books.
    """
    embeddings = OpenAIEmbeddings(
        api_key=OPENAI_API_KEY,
        base_url=OPENAI_BASE_URL,
    )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n# ", "\n## ", "\n### ", "\n", ". ", "! ", "? ", ",", " ", ""]
    )

    all_docs: List[LangchainDocument] = []

    # 1. JSON documents
    if Path(json_dir).exists():
        print(f"\n[1/2] Loading JSON documents: {json_dir}")
        json_docs = load_json_documents(json_dir)
        all_docs.extend(splitter.split_documents(json_docs))
        print(f"  JSON document chunks: {len(json_docs)}")
    else:
        print(f"[1/2] JSON directory does not exist, skipping: {json_dir}")

    # 2. DOCX documents
    if Path(docx_dir).exists():
        print(f"\n[2/2] Loading DOCX documents: {docx_dir}")
        docx_docs = load_docx_documents(docx_dir)
        all_docs.extend(docx_docs)
        print(f"  DOCX document chunks: {len(docx_docs)}")
    else:
        print(f"[2/2] DOCX directory does not exist, skipping: {docx_dir}")

    if not all_docs:
        raise ValueError("No documents found. Please check data directory paths.")

    print(f"\nTotal document chunks: {len(all_docs)}")
    print("Generating vector embeddings and writing to Chroma (may take a few minutes)...")

    vector_store = Chroma.from_documents(
        documents=all_docs,
        embedding=embeddings,
        persist_directory=persist_dir,
    )

    print(f"\nKnowledge base built and saved to: {persist_dir}")
    return vector_store


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_knowledge_base()
